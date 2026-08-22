"""Lightweight parsers: PDF (pdfplumber), Excel (pandas/openpyxl), CSV, raw text.

Every function is import-safe and lazy-loads heavy dependencies so the API
stays small and fast at startup (important on low-resource machines).
"""

from __future__ import annotations

import csv
import io
import re
from datetime import date
from pathlib import Path
from typing import Any

MAX_FUTURE_YEARS = 20  # allow planning a couple of decades ahead
MAX_PAST_YEARS = 200  # allow old family travel logs


def _year_is_sane(year: int) -> bool:
    """True when *year* is a plausible travel year (typos like 2918 are not)."""
    now = date.today().year
    return now - MAX_PAST_YEARS <= year <= now + MAX_FUTURE_YEARS


SOURCE_TYPES: dict[str, str] = {
    ".pdf": "pdf",
    ".xlsx": "xlsx",
    ".xls": "xlsx",
    ".csv": "csv",
    ".txt": "text",
    ".md": "text",
    ".markdown": "text",
    ".docx": "docx",
    ".png": "image",
    ".jpg": "image",
    ".jpeg": "image",
    ".webp": "image",
    ".gif": "image",
    ".bmp": "image",
    ".tiff": "image",
}

MAX_TEXT_CHARS = 60_000  # cap text sent toward the LLM


def detect_source_type(filename: str) -> str:
    """Return the human-readable source type for *filename*.

    Raises ``ValueError`` for unsupported extensions.
    """
    ext = Path(filename).suffix.lower()
    if ext not in SOURCE_TYPES:
        supported = ", ".join(sorted(SOURCE_TYPES))
        raise ValueError(f"Unsupported file type '{ext or '?'}'. Supported: {supported}.")
    return SOURCE_TYPES[ext]


def extract_content(filename: str, data: bytes) -> dict[str, Any]:
    """Return ``{'source_type', 'text', 'structured'}`` for a raw uploaded file.

    ``structured`` is a list of stop dicts when the source file contains
    recognisable columns (Location / Date / Notes / Lat / Lng); otherwise ``None``.
    """
    source_type = detect_source_type(filename)
    if source_type == "pdf":
        text, structured = _extract_pdf(data), None
    elif source_type == "xlsx":
        text, structured = _extract_excel(data)
    elif source_type == "csv":
        text, structured = _extract_csv(data)
    elif source_type == "docx":
        text, structured = _extract_docx(data), None
    elif source_type == "image":
        # No text to extract locally — the endpoint routes images to a vision
        # model (Omniroute auto/best-vision, Ollama vision fallback).
        text, structured = "", None
    else:
        text, structured = _decode_text(data), None
    return {
        "source_type": source_type,
        "text": text.strip()[:MAX_TEXT_CHARS],
        "structured": structured,
    }


# ---------------------------------------------------------------------------
# File-format extractors
# ---------------------------------------------------------------------------
def _extract_pdf(data: bytes) -> str:
    import pdfplumber  # lazy import for low startup cost

    pages: list[str] = []
    with pdfplumber.open(io.BytesIO(data)) as pdf:
        for page in pdf.pages:
            pages.append(page.extract_text() or "")
    return "\n\n".join(pages)


def _extract_docx(data: bytes) -> str:
    """Extract text from a .docx (Word) file."""
    from docx import Document  # python-docx, lazy import

    doc = Document(io.BytesIO(data))
    parts: list[str] = [p.text for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            if any(cells):
                parts.append(" | ".join(cells))
    return "\n".join(parts)


def _extract_excel(data: bytes) -> tuple[str, list[dict[str, Any]] | None]:
    try:
        import pandas as pd

        # header=None keeps title rows so we can detect the real header row.
        sheets = pd.read_excel(
            io.BytesIO(data), sheet_name=None, dtype=object, header=None
        )
        return _frames_to_content(list(sheets.values()))
    except ImportError:
        # pandas unavailable on this platform (Python 3.14+?) → fallback to openpyxl
        return _extract_excel_openpyxl(data)


def _extract_excel_openpyxl(data: bytes) -> tuple[str, list[dict[str, Any]] | None]:
    from openpyxl import load_workbook

    lines: list[str] = []
    all_stops: list[dict[str, Any]] = []
    workbook = load_workbook(io.BytesIO(data), read_only=True, data_only=True)
    try:
        for sheet in workbook.worksheets:
            grid = [list(row) for row in sheet.iter_rows(values_only=True)]
            if not grid:
                continue
            header_idx = _find_header_row(grid)
            headers = [str(c).strip() if c is not None else "" for c in grid[header_idx]]
            lines.append(" | ".join(headers))
            rows: list[dict[str, Any]] = []
            for raw in grid[header_idx + 1 :]:
                cells = ["" if c is None else str(c).strip() for c in raw]
                cells = (cells + [""] * len(headers))[: len(headers)]
                lines.append(" | ".join(cells))
                rows.append(dict(zip(headers, cells)))
            sheet_stops = _rows_to_stops(rows) or []
            all_stops.extend(sheet_stops)
    finally:
        workbook.close()
    return "\n".join(lines), all_stops or None


def _extract_csv(data: bytes) -> tuple[str, list[dict[str, Any]] | None]:
    try:
        import pandas as pd

        frame = pd.read_csv(io.BytesIO(data), dtype=object, header=None)
        return _frames_to_content([frame])
    except ImportError:
        pass
    text = _decode_text(data)
    try:
        reader = list(csv.DictReader(io.StringIO(text)))
    except csv.Error:
        reader = []
    # CSV with title rows: re-detect the header row from the raw lines.
    if reader:
        raw_lines = [l for l in text.splitlines() if l.strip()]
        grid = [l.split(",") for l in raw_lines]
        idx = _find_header_row(grid)
        if idx:
            headers = [h.strip() for h in grid[idx]]
            data_rows = [
                dict(zip(headers, (row + [""] * len(headers))[: len(headers)]))
                for row in grid[idx + 1 :]
            ]
            return text, _rows_to_stops(data_rows)
    return text, _rows_to_stops(reader)


def _frames_to_content(
    frames: list[Any],
) -> tuple[str, list[dict[str, Any]] | None]:
    """Convert pandas DataFrames to text + optional structured stops.

    Frames are read with ``header=None``, so the real header row is detected
    by scanning for known column names (handles title rows above headers).
    Each sheet is converted to stops separately (sheets may have different
    column layouts) and the results are merged.
    """
    lines: list[str] = []
    all_stops: list[dict[str, Any]] = []
    for frame in frames:
        frame = frame.fillna("")
        if frame.empty:
            continue
        grid = [[str(c).strip() if c is not None else "" for c in row.tolist()] for _, row in frame.iterrows()]
        header_idx = _find_header_row(grid)
        headers = grid[header_idx]
        lines.append(" | ".join(headers))
        rows: list[dict[str, Any]] = []
        for row in grid[header_idx + 1 :]:
            cells = (row + [""] * len(headers))[: len(headers)]
            lines.append(" | ".join(cells))
            rows.append(dict(zip(headers, cells)))
        sheet_stops = _rows_to_stops(rows) or []
        all_stops.extend(sheet_stops)
    return "\n".join(lines), all_stops or None


def _decode_text(data: bytes) -> str:
    for encoding in ("utf-8", "utf-16", "latin-1"):
        try:
            return data.decode(encoding)
        except (UnicodeDecodeError, UnicodeError):
            continue
    return data.decode("utf-8", errors="replace")


# ---------------------------------------------------------------------------
# Structured table detection (Location / Date / Notes / Lat / Lng columns)
# ---------------------------------------------------------------------------
COLUMN_ALIASES: dict[str, tuple[str, ...]] = {
    "location": ("location", "place", "city", "destination", "stop", "where", "name", "attraction", "venue", "site"),
    "exact_location": ("exact_location", "exact location", "full_location", "full location", "address", "region"),
    "date": ("date", "day", "when", "dates"),
    "start_date": ("start_date", "start date", "start", "from", "arrival", "arrive"),
    "end_date": ("end_date", "end date", "end", "to", "departure", "depart"),
    "notes": ("notes", "note", "description", "details", "activity", "activities", "comment", "comments", "summary", "what", "itinerary", "tips", "travel tips"),
    "category": ("category", "status", "trip", "classification"),
    "lat": ("lat", "latitude"),
    "lng": ("lng", "lon", "long", "longitude"),
    "country": ("country", "nation"),
}


def _norm_column(name: str) -> str:
    return re.sub(r"[^a-z0-9 ]", "", name.lower()).strip()


# Every known column word (normalized), used to detect the real header row in
# planner-style sheets that have title rows above the headers.
_KNOWN_COLUMN_WORDS = {
    _norm_column(a) for aliases in COLUMN_ALIASES.values() for a in aliases
}


def _looks_like_header(cell: Any) -> bool:
    """True when a cell text matches a known column name (exact or suffix)."""
    if cell is None:
        return False
    norm = _norm_column(str(cell))
    if not norm:
        return False
    return norm in _KNOWN_COLUMN_WORDS or any(
        norm.endswith(" " + k) for k in _KNOWN_COLUMN_WORDS
    )


def _find_header_row(rows: list[list[Any]]) -> int:
    """Index of the first row that looks like a column header.

    Planner-style spreadsheets often have title/subtitle rows above the real
    headers (e.g. "MASTER PLACES & POINTS OF INTEREST DATABASE"). We scan
    down for the first row with at least two recognizable column names.
    """
    for idx, row in enumerate(rows):
        matched = sum(1 for cell in row if _looks_like_header(cell))
        if matched >= 2:
            return idx
    return 0


def _split_multi_city(location: str) -> list[str]:
    """Split a "A, B, C, D" location cell into separate stop names.

    Cells with three or more comma-separated parts are almost always a list of
    distinct cities (e.g. "Venice, Florence, Pisa, Rome, Vatican"), so each part
    becomes its own stop so it can geocode and plot individually. Two-part names
    ("Heidelberg, Black Forest") are real compound place names and are kept as a
    single stop.
    """
    parts = [p.strip() for p in location.split(",") if p.strip()]
    if len(parts) >= 3:
        return parts
    return [location]


def _rows_to_stops(rows: list[dict[str, Any]]) -> list[dict[str, Any]] | None:
    """Convert table rows to stops when a location column exists, else ``None``."""
    if not rows:
        return None
    headers = list(rows[0].keys())

    mapped: dict[str, str | None] = {}
    year_header: str | None = None
    for role, aliases in COLUMN_ALIASES.items():
        mapped[role] = None
        for header in headers:
            norm = _norm_column(header)
            if any(norm == a or norm.endswith(" " + a) for a in aliases):
                mapped[role] = header
                break
        # Also detect a standalone "year" column for date supplementation.
        if role == "start_date" and mapped[role] is None:
            for header in headers:
                if _norm_column(header) == "year":
                    year_header = header
                    break

    if mapped["location"] is None and mapped["exact_location"] is None:
        return None

    def _supplement_year(found: str, row_val: Any) -> str:
        """If *found* is a partial MM-DD date, prepend the year from *row_val*.

        Years that are clearly typos (e.g. 2918 instead of 2018) are rejected,
        leaving the date as a partial — it will be dropped by ``_clean_iso`` in
        the API layer, producing an undated stop instead of a wildly wrong one.
        """
        if re.match(r"^\d{2}-\d{2}$", found):
            y = str(row_val).strip() if row_val is not None else ""
            y = re.sub(r"\D", "", y)[:4]
            if len(y) == 4 and y.isdigit() and _year_is_sane(int(y)):
                return f"{y}-{found}"
        return found

    stops: list[dict[str, Any]] = []
    for index, row in enumerate(rows):
        location = _clean(str(row.get(mapped["location"]) or "")) if mapped["location"] else ""
        exact = _clean(str(row.get(mapped["exact_location"]) or "")) if mapped["exact_location"] else ""
        if not location and not exact:
            continue

        start = _clean(str(row.get(mapped["start_date"]) or "")) if mapped["start_date"] else ""
        end = _clean(str(row.get(mapped["end_date"]) or "")) if mapped["end_date"] else ""
        legacy_date = _clean(str(row.get(mapped["date"]) or "")) if mapped["date"] else ""
        if not start:
            start = legacy_date
        if not end:
            end = start if (mapped["start_date"] or mapped["date"]) else ""

        notes = _clean(str(row.get(mapped["notes"]) or "")) if mapped["notes"] else ""

        # Fallback: extract dates from the notes / itinerary text column when
        # no dedicated date columns exist.
        if not start and notes:
            found = _find_date(notes)
            if found:
                start = _supplement_year(found, row.get(year_header) if year_header else None)
        if not end and start:
            end = start

        category = _clean(str(row.get(mapped["category"]) or "")) if mapped["category"] else ""
        category = {"past": "Past", "current": "Current", "upcoming": "Upcoming"}.get(
            category.lower(), ""
        )
        if not category:
            category = classify_category(start, end)

        lat = _to_float(row.get(mapped["lat"])) if mapped["lat"] else None
        lng = _to_float(row.get(mapped["lng"])) if mapped["lng"] else None

        # Context hint: when the sheet has a Country column but no explicit
        # exact_location column, append the country to bare city names so they
        # geocode to the right place — e.g. "Sun City" → "Sun City, South
        # Africa" (not Sun City, Arizona) and "Nara" → "Nara, Japan" (not the
        # US National Archives, which matches the "NARA" acronym).
        country_hint = ""
        if mapped["country"] and mapped["exact_location"] is None:
            raw_country = _clean(str(row.get(mapped["country"]) or ""))
            if raw_country and raw_country.lower() not in (
                "—", "-", "n/a", "na", "none", "unknown", "tbd"
            ):
                country_hint = raw_country

        # Multi-city cells become one stop per city so each one geocodes and
        # plots on the map (e.g. "Venice, Florence, Pisa" → 3 stops).
        names = _split_multi_city(location or exact)
        for name in names:
            if len(names) == 1 and exact:
                # A dedicated exact_location column wins for single-city rows
                # (previously it was discarded in favour of the bare name).
                exact_loc = exact
            elif country_hint:
                exact_loc = f"{name}, {country_hint}"
            else:
                exact_loc = name
            stops.append(
                {
                    "order": len(stops) + 1,
                    "location": name,
                    "exact_location": exact_loc,
                    "start_date": start,
                    "end_date": end,
                    "category": category,
                    "notes": notes,
                    "lat": lat,
                    "lng": lng,
                }
            )
    return stops


def _clean(value: str) -> str:
    value = re.sub(r"\s+", " ", value).strip()
    if value.lower() in ("nan", "none", "nat", "null"):
        return ""
    return value


def _to_float(value: Any) -> float | None:
    if value is None or value == "":
        return None
    try:
        return float(value)
    except (TypeError, ValueError):
        return None


# ---------------------------------------------------------------------------
# Local fallback: extract likely locations without any LLM.
# ---------------------------------------------------------------------------
_MONTHS: dict[str, int] = {
    "jan": 1, "feb": 2, "mar": 3, "apr": 4, "may": 5, "jun": 6,
    "jul": 7, "aug": 8, "sep": 9, "oct": 10, "nov": 11, "dec": 12,
}

_DATE_PATTERNS = [
    re.compile(r"\b(\d{4})-(\d{1,2})-(\d{1,2})\b"),
    re.compile(r"\b(\d{1,2})[/.](\d{1,2})[/.](\d{2,4})\b"),
    re.compile(
        r"\b(Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-z]*\.?\s+"
        r"(\d{1,2})(?:st|nd|rd|th)?(?:,?\s+(\d{4}))?\b",
        re.I,
    ),
]


def _find_date(line: str) -> str:
    """Return the first date found in *line* as ISO ``YYYY-MM-DD`` or ``MM-DD``."""
    for pattern in _DATE_PATTERNS[:2]:
        m = pattern.search(line)
        if not m:
            continue
        first, second, third = m.groups()
        if len(third) == 2:
            third = ("20" if int(third) <= 35 else "19") + third
        year, month, day = int(third), int(second), int(first)
        if 1 <= month <= 12 and 1 <= day <= 31:
            return f"{year:04d}-{month:02d}-{day:02d}"
    m = _DATE_PATTERNS[2].search(line)
    if m:
        month = _MONTHS[m.group(1)[:3].lower()]
        day = int(m.group(2))
        year = m.group(3)
        base = f"{month:02d}-{day:02d}"
        return f"{year}-{base}" if year else base
    return ""


def _strip_noise(line: str) -> str:
    """Remove date/time tokens from *line* so the remaining text is the location name."""
    s = line
    for pattern in _DATE_PATTERNS:
        s = pattern.sub(" ", s)
    s = re.sub(r"\b\d{1,2}:\d{2}\s*(?:am|pm)?\b", " ", s, flags=re.I)
    s = re.sub(r"\b\d{1,2}\s*(?:am|pm)\b", " ", s, flags=re.I)
    return s


def heuristic_parse(text: str) -> list[dict[str, Any]]:
    """Very light, dependency-free location extraction (LLM offline fallback).

    The heuristic is deliberately loose — it favours returning *some* candidates
    over returning nothing, because the user can always edit the table.
    """
    stops: list[dict[str, Any]] = []
    seen: set[str] = set()
    skip_words = {
        "itinerary", "day", "agenda", "schedule", "travel plan", "trip",
        "route", "details", "notes",
    }

    for raw in text.splitlines():
        line = raw.strip()
        if not line:
            continue

        date = _find_date(line)
        candidate = _strip_noise(line)

        # Remove "Day N …" prefix, numbered-list prefixes, and bullet markers
        candidate = re.sub(r"^(?:day\s*\d+\s*[-–:.,)]*\s*)", "", candidate, flags=re.I)
        candidate = re.sub(r"^\s*(?:\d{1,3}\s*[-–.:)]?\s*|[-•*]\s+)", "", candidate)
        candidate = re.sub(r"[-–]\s*\d{1,2}\s*$", "", candidate)
        candidate = re.sub(r"\s+", " ", candidate).strip(" -–.:,;/")

        if len(candidate) < 3 or len(candidate) > 90:
            continue
        if candidate.lower() in skip_words:
            continue
        if re.fullmatch(r"[\d\s%$€£,./-]+", candidate):
            continue

        key = candidate.lower()
        if key in seen:
            continue
        seen.add(key)

        stops.append(
            {
                "order": len(stops) + 1,
                "location": candidate,
                "exact_location": candidate,
                "start_date": date,
                "end_date": date,
                "category": classify_category(date, date) if date else "",
                "notes": "",
                "lat": None,
                "lng": None,
            }
        )
    return stops


# ---------------------------------------------------------------------------
# Date classification + travel feasibility (post-processing)
# ---------------------------------------------------------------------------
def classify_category(start_date: str, end_date: str, today: date | None = None) -> str:
    """Classify a stop as ``"Past"`` / ``"Current"`` / ``"Upcoming"`` (or ``""``).

    A stop is *Current* when today falls inside its date range (or on its single
    start date); *Past* when it ended before today; *Upcoming* when it starts
    after today.
    """

    def to_date(value: str) -> date | None:
        try:
            d = date.fromisoformat(value.strip())
        except (ValueError, AttributeError):
            return None
        # Reject typo years (e.g. 2918): an impossible date has no category.
        return d if _year_is_sane(d.year) else None

    start, end = to_date(start_date), to_date(end_date)
    if not start and not end:
        return ""
    today = today or date.today()
    if end:
        if end < today:
            return "Past"
        if start and start > today:
            return "Upcoming"
        return "Current"
    if start < today:
        return "Past"
    if start > today:
        return "Upcoming"
    return "Current"


def _parse_iso_date(value: Any) -> date | None:
    try:
        return date.fromisoformat(str(value).strip()[:10])
    except ValueError:
        return None


def feasibility_check(
    stops: list[dict[str, Any]],
    speed_limit_kmh: float = 900.0,
    min_travel_hours: float = 8.0,
) -> list[dict[str, Any]]:
    """Flag stops whose implied travel speed from the previous stop is impossible.

    For each adjacent pair with coordinates and dates, the geodesic distance
    (``geopy.distance``) divided by the time gap gives an implied speed. If that
    speed exceeds realistic human travel limits (``speed_limit_kmh``), the arriving
    stop is marked ``is_ambiguous: True`` — the wrong region (e.g. "Hamilton, USA"
    instead of "Hamilton, NZ") was probably picked.

    Returns one entry per stop: ``{order, location, exact_location, is_ambiguous,
    warning, candidates}`` (``candidates`` is filled in later by the API layer).
    """
    from geopy.distance import geodesic  # lazy import keeps startup light

    entries: list[dict[str, Any]] = []
    for i, stop in enumerate(stops):
        name = stop.get("location") or ""
        exact = stop.get("exact_location") or name
        entry: dict[str, Any] = {
            "order": stop.get("order", i + 1),
            "location": name,
            "exact_location": exact,
            "is_ambiguous": False,
            "warning": None,
            "candidates": [],
        }
        entries.append(entry)
        if i == 0:
            continue

        prev = stops[i - 1]
        if (
            prev.get("lat") is None
            or prev.get("lng") is None
            or stop.get("lat") is None
            or stop.get("lng") is None
        ):
            continue
        prev_dt = _parse_iso_date(prev.get("end_date") or prev.get("start_date"))
        curr_dt = _parse_iso_date(stop.get("start_date") or stop.get("end_date"))
        if prev_dt is None or curr_dt is None:
            continue

        dist_km = geodesic(
            (float(prev["lat"]), float(prev["lng"])),
            (float(stop["lat"]), float(stop["lng"])),
        ).km
        hours = (curr_dt - prev_dt).total_seconds() / 3600.0
        speed = dist_km / max(hours, min_travel_hours)

        if speed > speed_limit_kmh:
            prev_name = prev.get("exact_location") or prev.get("location") or "previous stop"
            entry["is_ambiguous"] = True
            entry["warning"] = (
                f"Implied travel speed {speed:.0f} km/h from “{prev_name}” to “{exact}” "
                f"({dist_km:,.0f} km in {max(hours, min_travel_hours):.1f}h) is physically "
                f"implausible — the wrong region may have been picked for “{name}”. "
                f"Please confirm the correct country/region below."
            )
    return entries